"""
clauses/agreement_renderer.py — Entry point for Agreement Rendering Engine
===========================================================================
Re-exports:
  - generate_preview_html(data) -> HTML preview for frontend
  - generate_docx(data, template_filename, output_path) -> DOCX file export
"""

import sys, os, re
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from .formatters import (
    num_to_words, _safe_int, combine_name_prefix_once, format_careof,
    format_age, format_ordinal_date, clean_text, format_rent_increase,
    _is_leave_license, _should_skip_clause, _add_runs_with_bold
)
from .evaluator import (
    _PKEY_TO_CANONICAL, _CANONICAL_ALIASES, _resolve_value,
    _build_field_map, _evaluate_clause, _substitute_fields
)
from .html_renderer import generate_preview_html, _get_preamble_paragraphs
from .leave_license import CLAUSES as LEAVE_LICENSE_CLAUSES
from .simple_rental import CLAUSES as SIMPLE_RENTAL_CLAUSES

def generate_docx(data, template_filename=None, output_path=None):
    """Build the rendered agreement as a Word document and save it to output_path."""
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    if not output_path:
        raise ValueError("output_path is required for DOCX generation")

    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    agreement_type = _resolve_value(data, "agreement_type") or "simple_rental"
    is_leave_license = _is_leave_license(agreement_type)
    clause_defs = LEAVE_LICENSE_CLAUSES if is_leave_license else SIMPLE_RENTAL_CLAUSES
    field_map = _build_field_map(data)

    document = Document()
    section = document.sections[0]
    # Match the supplied rental-agreement reference: A4 with a compact top
    # margin and a readable legal-document body.
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.39)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = document.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    normal.font.size = Pt(13)
    normal.paragraph_format.space_after = Pt(0)

    # Bold only meaningful substituted values. Short values such as "his" or
    # relationship abbreviations must not bold matching words in legal prose.
    raw_targets = {
        "Owner", "Owners", "Tenant", "Tenants",
        "Licensor", "Licensors", "Licensee", "Licensees",
        "OWNER", "OWNERS", "TENANT", "TENANTS",
        "LICENSOR", "LICENSORS", "LICENSEE", "LICENSEES"
    }
    pen_val = _resolve_value(data, "penalty_deduction")
    if pen_val:
        pen_str = str(pen_val).strip()
        raw_targets.add(pen_str)
        if not pen_str.lower().endswith("days"):
            raw_targets.add(f"{pen_str} days")
            raw_targets.add(f"{pen_str} DAYS")

    for value in field_map.values():
        if value:
            v_str = str(value).strip()
            if len(v_str) >= 4:
                raw_targets.add(v_str)
                if '\n' in v_str:
                    for line_part in v_str.split('\n'):
                        line_clean = line_part.strip()
                        if len(line_clean) >= 4:
                            raw_targets.add(line_clean)

    bold_targets = sorted(raw_targets, key=len, reverse=True)
    party_age_targets = {
        f"{data.get(f'{party}{index}_age', '')} years"
        for party in ('owner', 'tenant')
        for index in range(1, 7)
        if str(data.get(f'{party}{index}_age', '')).strip()
    }

    def style_run(run, bold=False, underline=False, size=13):
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        run.font.size = Pt(size)
        run.bold = bold
        run.underline = underline
        return run

    def add_rich_text(paragraph, text, extra_bold=()):
        targets = sorted(set(raw_targets).union(extra_bold), key=len, reverse=True)
        _add_runs_with_bold(paragraph, text, targets)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('LEAVE AND LICENSE AGREEMENT' if is_leave_license else 'RENTAL AGREEMENT')
    title_run.bold = True
    title_run.underline = True
    style_run(title_run, bold=True, underline=True, size=14)
    title.paragraph_format.space_after = Pt(18)

    for text in _get_preamble_paragraphs(data, field_map, agreement_type):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(12)
        if text in ('BETWEEN', 'AND'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_run(paragraph.add_run(text), bold=True)
        elif text.startswith('NOW THEREFORE') or text.startswith('NOW THIS AGREEMENT'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(14)
            paragraph.paragraph_format.space_after = Pt(14)
            style_run(paragraph.add_run(text), bold=True, underline=True)
        else:
            for index, line in enumerate(text.split('\n')):
                if index:
                    paragraph.add_run().add_break()
                add_rich_text(paragraph, line, extra_bold=party_age_targets)

            # The reference uses bold party-role statements immediately after
            # the party particulars, while ordinary legal text stays regular.
            if text.startswith('(Here in after called') or text.startswith('Hereinafter called'):
                for run in paragraph.runs:
                    run.bold = True

    clause_number = 1
    for clause in clause_defs:
        if _should_skip_clause(clause, agreement_type):
            continue
        evaluated_text = _evaluate_clause(clause, data)
        if evaluated_text is None:
            continue
        text = _substitute_fields(evaluated_text, field_map)

        if clause.get('is_header'):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(16)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            style_run(paragraph.add_run(text), bold=True, underline=True)
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        style_run(paragraph.add_run(f'{clause_number}. '), bold=True)
        if clause.get('title'):
            style_run(paragraph.add_run(f"{clause['title']}: "), bold=True)
        add_rich_text(paragraph, text)
        clause_number += 1

    # ── CONCLUDING / SIGNATURE BLOCK (DOCX) ──────────────────────────────────
    owner_count = _safe_int(_resolve_value(data, "owner_count") or "1")
    tenant_count = _safe_int(_resolve_value(data, "tenant_count") or "1")

    if is_leave_license:
        owner_role_label = "LICENSORS" if owner_count > 1 else "LICENSOR"
        tenant_role_label = "LICENSEES" if tenant_count > 1 else "LICENSEE"
    else:
        owner_role_label = "OWNERS" if owner_count > 1 else "OWNER"
        tenant_role_label = "TENANTS" if tenant_count > 1 else "TENANT"

    # Description of Premises Header
    p_desc_hdr = document.add_paragraph()
    p_desc_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc_hdr.paragraph_format.space_before = Pt(24)
    p_desc_hdr.paragraph_format.space_after = Pt(12)
    style_run(p_desc_hdr.add_run("DESCRIPTION OF THE SAID PREMISES:"), bold=True, underline=True)

    # Description of Premises Body
    prop_address = field_map.get("{property_address}", "")
    p_desc = document.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(12)
    p_desc.add_run("All that consisting of premises, at ")
    style_run(p_desc.add_run(f"{prop_address}."), bold=True)

    # In Witness Clause
    agr_date = field_map.get("{agreement_date}", "")
    p_wit = document.add_paragraph()
    p_wit.paragraph_format.space_after = Pt(18)
    p_wit.add_run("IN WITNESS, WHERE OF THE PARTIES TO THIS AGREEMENT HAVE SIGNED HEREUNDER ON THE AFORESAID DATE ")
    style_run(p_wit.add_run(f"{agr_date}"), bold=True)

    # Owners / Licensors
    p_own_hdr = document.add_paragraph()
    p_own_hdr.paragraph_format.space_before = Pt(12)
    p_own_hdr.paragraph_format.space_after = Pt(12)
    style_run(p_own_hdr.add_run(f"{owner_role_label}:"), bold=True)

    for i in range(1, owner_count + 1):
        o_pfx = _resolve_value(data, f"owner{i}_prefix") or ""
        o_name_raw = _resolve_value(data, f"owner{i}_name") or ""
        o_name = combine_name_prefix_once(o_pfx, o_name_raw)

        p_o = document.add_paragraph()
        p_o.paragraph_format.space_after = Pt(14)
        p_o.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT)
        p_o.add_run("Name: ")
        style_run(p_o.add_run(o_name), bold=True)
        style_run(p_o.add_run("\tSignature"), bold=True)

    # Tenants / Licensees
    p_ten_hdr = document.add_paragraph()
    p_ten_hdr.paragraph_format.space_before = Pt(14)
    p_ten_hdr.paragraph_format.space_after = Pt(12)
    style_run(p_ten_hdr.add_run(f"{tenant_role_label}:"), bold=True)

    for i in range(1, tenant_count + 1):
        t_pfx = _resolve_value(data, f"tenant{i}_prefix") or ""
        t_name_raw = _resolve_value(data, f"tenant{i}_name") or ""
        t_name = combine_name_prefix_once(t_pfx, t_name_raw)

        p_t = document.add_paragraph()
        p_t.paragraph_format.space_after = Pt(14)
        p_t.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT)
        p_t.add_run("Name: ")
        style_run(p_t.add_run(t_name), bold=True)
        style_run(p_t.add_run("\tSignature"), bold=True)

    # Witness Section
    witness_header = "IN THE PRESENCE OF WITNESS" if is_leave_license else "IN THE PRESENCE OF WITNESS:"
    p_w_hdr1 = document.add_paragraph()
    p_w_hdr1.paragraph_format.space_before = Pt(18)
    p_w_hdr1.paragraph_format.space_after = Pt(4)
    style_run(p_w_hdr1.add_run(witness_header), bold=True)

    p_w_hdr2 = document.add_paragraph()
    p_w_hdr2.paragraph_format.space_after = Pt(12)
    style_run(p_w_hdr2.add_run("WITNESS"), bold=True)

    p_w_line = document.add_paragraph()
    p_w_line.paragraph_format.space_after = Pt(18)
    p_w_line.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT)
    p_w_line.add_run("Name:")
    style_run(p_w_line.add_run("\tSignature"), bold=True)

    # Annexure (DOCX)
    annexure = str(_resolve_value(data, "annexure") or "").strip()
    if annexure:
        p_ann_hdr = document.add_paragraph()
        p_ann_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ann_hdr.paragraph_format.space_before = Pt(36)
        p_ann_hdr.paragraph_format.space_after = Pt(14)
        style_run(p_ann_hdr.add_run("ANNEXURE"), bold=True, underline=True)

        p_ann_body = document.add_paragraph()
        p_ann_body.paragraph_format.space_after = Pt(12)
        annexure_lines = [line.strip().upper() for line in annexure.splitlines()]
        for idx, a_line in enumerate(annexure_lines):
            if idx > 0:
                p_ann_body.add_run().add_break()
            style_run(p_ann_body.add_run(a_line), bold=False)

    document.save(output_path)
    return output_path
